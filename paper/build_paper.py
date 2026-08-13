from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Durable_Goals_IEEE_Writeup.md"
OUTPUT = ROOT / "Durable_Goals_IEEE_Writeup.docx"

INK = "1F2430"
MUTED = "596579"
RULE = "B7BEC8"
LIGHT = "EEF1F4"
ACCENT = "B34B22"
BODY_FONT = "Times New Roman"
MONO_FONT = "Courier New"


def set_run_font(run, name=BODY_FONT, size=9.5, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_keep(paragraph, *, next_paragraph=False, together=False):
    ppr = paragraph._p.get_or_add_pPr()
    if next_paragraph:
        ppr.append(OxmlElement("w:keepNext"))
    if together:
        ppr.append(OxmlElement("w:keepLines"))
    ppr.append(OxmlElement("w:widowControl"))


def set_section_geometry(section, columns=1):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(columns))
    cols.set(qn("w:space"), "360")  # 0.25 in
    cols.set(qn("w:equalWidth"), "1")


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = True
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=7.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(3.2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Inches(0.14)

    for name, size, before, after, align, bold, italic in (
        ("Title", 22, 0, 7, WD_ALIGN_PARAGRAPH.CENTER, False, False),
        ("Subtitle", 10.5, 0, 2, WD_ALIGN_PARAGRAPH.CENTER, False, False),
        ("Heading 1", 10, 8, 3, WD_ALIGN_PARAGRAPH.CENTER, True, False),
        ("Heading 2", 9.5, 6, 2, WD_ALIGN_PARAGRAPH.LEFT, False, True),
        ("Heading 3", 9.5, 5, 2, WD_ALIGN_PARAGRAPH.LEFT, True, False),
        ("Caption", 7.8, 2, 5, WD_ALIGN_PARAGRAPH.JUSTIFY, False, False),
    ):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor.from_string(INK if name != "Caption" else MUTED)
        p = style.paragraph_format
        p.alignment = align
        p.space_before = Pt(before)
        p.space_after = Pt(after)
        p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.keep_with_next = name != "Caption"
        p.keep_together = True

        # Built-in Word title styles can carry theme borders. The IEEE override
        # uses typography only, so remove any inherited paragraph border.
        ppr = style._element.get_or_add_pPr()
        border = ppr.find(qn("w:pBdr"))
        if border is not None:
            ppr.remove(border)

    styles["Heading 1"].paragraph_format.page_break_before = False


def set_language(doc):
    styles = doc.styles
    for style in styles:
        if style.type != 1:  # paragraph style
            continue
        rpr = style._element.get_or_add_rPr()
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:val"), "en-US")


def add_hyperlink(paragraph, text, url, *, size=7.7):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E5AA8")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(round(size * 2)))
    for child in (rfonts, color, underline, sz):
        rpr.append(child)
    new_run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


URL_RE = re.compile(r"https?://[^\s]+")


def add_text_with_links(paragraph, text, *, size=9.5, bold=False, italic=False, color=INK):
    position = 0
    for match in URL_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size, bold=bold, italic=italic, color=color)
        url_text = match.group(0)
        trailing = ""
        while url_text and url_text[-1] in ".,)":
            trailing = url_text[-1] + trailing
            url_text = url_text[:-1]
        add_hyperlink(paragraph, url_text, url_text, size=size)
        if trailing:
            run = paragraph.add_run(trailing)
            set_run_font(run, size=size, bold=bold, italic=italic, color=color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def add_body_paragraph(doc, text, *, reference=False):
    p = doc.add_paragraph()
    if reference:
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2.4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        add_text_with_links(p, text, size=7.7)
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.14)
        p.paragraph_format.space_after = Pt(3.2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        add_text_with_links(p, text, size=9.5)
    set_paragraph_keep(p, together=True)
    return p


def create_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = 710
    num_id = 710
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "50")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    level.extend([start, num_fmt, lvl_text, lvl_jc, ppr])
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)
    add_text_with_links(p, text, size=9.5)
    set_paragraph_keep(p, together=True)
    return p


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=50, start=60, bottom=50, end=60):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "60")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tc_w = tcpr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tcpr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), RULE)


def add_table(doc, rows):
    columns = len(rows[0])
    if rows[0][0] == "RECORD":
        widths = [1500, 1200, 2060]
    elif rows[0][0] == "TEST MODULE":
        widths = [1300, 900, 2560]
    else:
        widths = [1900, 2860]
    table = doc.add_table(rows=len(rows), cols=columns)
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(repeat)
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ridx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (ridx == 0 or cidx == 1) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=7.2, bold=(ridx == 0), color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, path, width, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    set_paragraph_keep(p, next_paragraph=True, together=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=7.8, color=MUTED)
    set_paragraph_keep(p, together=True)
    return p


def add_equation(doc, text, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_together = True
    run = p.add_run(f"{text}    ({number})")
    set_run_font(run, name="Cambria Math", size=9.5, italic=True, color=INK)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.05)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    ppr.append(shd)
    run = p.add_run(text)
    set_run_font(run, name=MONO_FONT, size=6.5, color=INK)
    set_paragraph_keep(p, together=True)


def add_front_matter(doc, title, author_lines, abstract, index_terms):
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=22, color=INK)
    set_paragraph_keep(p, next_paragraph=True, together=True)

    for idx, line in enumerate(author_lines):
        p = doc.add_paragraph(style="Subtitle")
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx == 0:
            set_run_font(p.add_run(line), size=11, color=INK)
        else:
            set_run_font(p.add_run(line), size=8.5, italic=(idx == 1), color=MUTED)
        set_paragraph_keep(p, next_paragraph=True, together=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(p.add_run("Abstract—"), size=8.6, bold=True, italic=True, color=INK)
    set_run_font(p.add_run(abstract), size=8.6, color=INK)
    set_paragraph_keep(p, together=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(p.add_run("Index Terms—"), size=8.6, bold=True, italic=True, color=INK)
    set_run_font(p.add_run(index_terms), size=8.6, color=INK)
    set_paragraph_keep(p, together=True)


def parse_source():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0][2:].strip()
    idx = 2
    author_lines = []
    while idx < len(lines) and not lines[idx].startswith("ABSTRACT:"):
        if lines[idx].strip():
            author_lines.append(lines[idx].strip())
        idx += 1
    abstract = lines[idx].split("ABSTRACT:", 1)[1].strip()
    idx += 2
    index_terms = lines[idx].split("INDEX TERMS:", 1)[1].strip()
    idx += 1
    return title, author_lines, abstract, index_terms, lines[idx:]


def build():
    title, author_lines, abstract, index_terms, lines = parse_source()
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author_lines[0]
    doc.core_properties.subject = "IEEE-style technical report for Durable Goals 0.1.1"
    doc.core_properties.keywords = "AI agents, durable goals, provenance, workflow DAG"
    configure_styles(doc)
    set_language(doc)
    for section in doc.sections:
        set_section_geometry(section, columns=1)
    add_page_number(doc.sections[0])

    add_front_matter(doc, title, author_lines, abstract, index_terms)
    bullet_num_id = create_bullet_numbering(doc)

    body_started = False
    references = False
    equation_number = 0
    paragraph_buffer = []
    table_caption_pending = False
    full_width_figure = False

    def flush_paragraph():
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer)
        add_body_paragraph(doc, text, reference=references)
        paragraph_buffer = []

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            if not body_started:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                set_section_geometry(section, columns=2)
                body_started = True
            heading = stripped[3:]
            references = heading.endswith("REFERENCES")
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.first_line_indent = None
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(heading.upper())
            set_run_font(run, size=10, bold=True, color=INK)
            set_paragraph_keep(p, next_paragraph=True, together=True)
            index += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.first_line_indent = None
            run = p.add_run(stripped[4:])
            set_run_font(run, size=9.5, italic=True, color=INK)
            set_paragraph_keep(p, next_paragraph=True, together=True)
            index += 1
            continue

        figure_match = re.fullmatch(r"\[\[FIGURE:(.+)\|([0-9.]+)\]\]", stripped)
        if figure_match:
            flush_paragraph()
            figure_path = ROOT / figure_match.group(1)
            width = float(figure_match.group(2))
            if body_started:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                set_section_geometry(section, columns=1)
                full_width_figure = True
            alt = (
                "Diagram of Durable Goals resolution and activation semantics"
                if "resolution" in figure_path.name
                else "Diagram of optional DAG prompt eligibility and external harness execution"
            )
            add_figure(doc, figure_path, width, alt)
            index += 1
            continue

        caption_match = re.fullmatch(r"\[\[CAPTION:(.+)\]\]", stripped)
        if caption_match:
            flush_paragraph()
            add_caption(doc, caption_match.group(1))
            if full_width_figure:
                section = doc.add_section(WD_SECTION.CONTINUOUS)
                set_section_geometry(section, columns=2)
                full_width_figure = False
            index += 1
            continue

        equation_match = re.fullmatch(r"\[\[EQUATION:(.+)\]\]", stripped)
        if equation_match:
            flush_paragraph()
            equation_number += 1
            add_equation(doc, equation_match.group(1), equation_number)
            index += 1
            continue

        code_match = re.fullmatch(r"\[\[CODE:(.+)\]\]", stripped)
        if code_match:
            flush_paragraph()
            add_code(doc, code_match.group(1))
            index += 1
            continue

        if stripped.startswith("TABLE "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped)
            set_run_font(run, size=7.8, bold=True, color=INK)
            set_paragraph_keep(p, next_paragraph=True, together=True)
            table_caption_pending = True
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            raw_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                raw_rows.append(lines[index].strip())
                index += 1
            parsed = []
            for raw in raw_rows:
                cells = [cell.strip() for cell in raw.strip("|").split("|")]
                if all(re.fullmatch(r":?-+:?", cell) for cell in cells):
                    continue
                parsed.append(cells)
            add_table(doc, parsed)
            table_caption_pending = False
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_bullet(doc, stripped[2:], bullet_num_id)
            index += 1
            continue

        if references and re.match(r"^\[\d+\]", stripped):
            flush_paragraph()
            add_body_paragraph(doc, stripped, reference=True)
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()

    # Apply page-number field and geometry consistently to every generated section.
    for section in doc.sections:
        if section.footer.paragraphs:
            section.footer.paragraphs[0].paragraph_format.space_after = Pt(0)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
